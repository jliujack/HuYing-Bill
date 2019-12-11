
import docx
import datetime
import time
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph  import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from handleBill import util
import handleBill.assist as assist

CONFIG_COL, WORK_COL = util.CONFIG_COL, util.WORK_COL

MONTH_EN = ("Zero", "January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December" )

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def docx_to_list(docx):
    result = []
    for block in iter_block_items(docx):
        if isinstance(block, Paragraph):
            result.append(block.text)
        elif isinstance(block, Table):
            temp = []
            for row in block.rows:
                cell_list = []
                for cell in row.cells:
                    for elment in iter_block_items(cell):
                        cell_list.append(elment.text)
                temp.append(cell_list)
            result.append(temp)
    return result


def print_docx(docx):
    for block in iter_block_items(docx):
        if isinstance(block, Paragraph):
            print(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for elment in iter_block_items(cell):
                        print(elment.text+'\t', end="")
                print('\n')


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def replaceCell(cell, text):
  cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
  run = None
  if len(cell.paragraphs[0].runs) > 0:
    run = cell.paragraphs[0].runs[0]
    run.text = text
    for i in range(1, len(cell.paragraphs[0].runs)):
      cell.paragraphs[0].runs[i].text = ""
  else:
    run = cell.paragraphs[0].add_run(text)
  run.font.size = Pt(10.5)
  return run

def replaceTable2(table2, caseWorkItems):
  orgDate = table2.cell(2, 0).text
  for i in range(len(table2.rows) - 1, 0, -1):
    remove_row(table2, table2.rows[i])
  for i in range(0, len(caseWorkItems)):
    table2.add_row()
    cells = table2.add_row().cells
    for j in range(4):
      text = str(caseWorkItems[i][j])
      if j == 0 and text != "":
        date = datetime.datetime.strptime(text,'%Y-%m-%d')
        if len(orgDate) > 8:
          text = date.strftime('%d/%m/%Y')
        else:
          text = date.strftime('%d/%m/%y')
      run = replaceCell(cells[j], text) 
      run.font.name='Times New Roman'
      r = run._element
      r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
      cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
      if j > 1:
        cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
  table2.add_row()

def replaceCell3(cell, text):
  cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  run = None
  if len(cell.paragraphs[0].runs) > 0:
    run = cell.paragraphs[0].runs[0]
    run.text = text
    for i in range(1, len(cell.paragraphs[0].runs)):
      cell.paragraphs[0].runs[i].text = ""
  else:
    run = cell.paragraphs[0].add_run(text)
  run.font.size = Pt(10.5)
  return run

def replaceTable3(table3, billStatistics, total):
  orgDate = table3.cell(2, 0).text
  for i in range(len(table3.rows) - 1, 0, -1):
    remove_row(table3, table3.rows[i])
  for i in range(0, len(billStatistics)):
    table3.add_row()
    cells = table3.add_row().cells
    for j in range(4):
      text = str(billStatistics[i][j])
      if j == 1:
        text = "{:,}".format(round(float(billStatistics[i][j]), 1))
      if j > 1:
        text = "{:,}".format(round(float(billStatistics[i][j]), 2))
      run = replaceCell(cells[j], text) 
      run.font.name='Times New Roman'
      r = run._element
      r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
      cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
  table3.add_row()

  cells = table3.add_row().cells
  run = replaceCell3(cells[-2], "Total/小计:")
  run.bold = True
  r = run._element
  # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
  cells[-2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

  run = replaceCell3(cells[-1], total)
  run.bold = True
  run.font.name='Times New Roman'
  cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def replaceNormalWord(targetPath, caseWorkItems, billStatistics, total, date, config):
  date = datetime.datetime.strptime(date,'%Y-%m-%d')
  time.sleep(1)
  total = "{:,}".format(round(total,2))
  try:
    docx_file = docx.Document(targetPath)

    # 替换表格1的总结性内容
    table1 = docx_file.tables[0]
    replaceCell(table1.cell(0, 1), date.strftime('%d/%m/%Y'))
    replaceCell(table1.cell(0, 3), str(config[CONFIG_COL.contractNum .value]) + "-0001/" + str(config[CONFIG_COL.versionNum.value]))
    replaceCell(table1.cell(2, 3), MONTH_EN[date.month] + " " + str(date.year))
    replaceCell(table1.cell(3, 1), config[CONFIG_COL.currency.value] + " " + total)
    replaceCell(table1.cell(4, 1), config[CONFIG_COL.currency.value] + " " + total)

    #替换表格2的收费明细
    table2 = docx_file.tables[2]
    replaceTable2(table2, caseWorkItems)

    #替换表格3的统计数据
    table3 = docx_file.tables[3]
    replaceTable3(table3, billStatistics, total)
    
    # 必须要保存才能修改生效
    docx_file.save(targetPath)
  except Exception as e:
    print(e)